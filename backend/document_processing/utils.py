import pytesseract
import os
from pdf2image import convert_from_path
from docx import Document
from django.conf import settings
import re
import pdfplumber
import docx
from .groq_api import query_groq_resume_extraction
import json


def extract_sections(text):
    doc = nlp(text)
    entities = {}
    for ent in doc.ents:
        if ent.label_ in entities:
            entities[ent.label_].append(ent.text)
        else:
            entities[ent.label_] = [ent.text]
    return entities


def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file."""
    with pdfplumber.open(pdf_path) as pdf:
        return "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file."""
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def clean_groq_response(response_text):
    """Extract JSON part from the Groq response using regex."""
    print("The response is ", response_text)
    match = re.search(r"```json\n(.*?)\n```", response_text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))  # Convert string to JSON
        except json.JSONDecodeError:
            return {"error": "Invalid JSON extracted"}
    return {"error": "No JSON found in response"}

def extract_personal_info(data):
    """Extracts Name, Email, and Phone from the structured JSON response."""
    return {
        "name": data.get("Personal Information", {}).get("Name", "Not Found"),
        "email": data.get("Personal Information", {}).get("Email", "Not Found"),
        "phone": data.get("Personal Information", {}).get("Phone", "Not Found")
    }

def extract_education(data):
    """Extracts education details and formats them according to the model."""
    education_list = []
    for edu in data.get("Education", []):
        education_list.append({
            "degree": edu.get("Degree", "Unknown"),
            "university": edu.get("University", "Unknown"),
            "field": edu.get("Field", "Unknown"),
            "start_year": str(edu.get("Start Year", "0000")),
            "end_year": str(edu.get("End Year", "")) if edu.get("End Year") else None
        })
    return education_list

def extract_skills(data):
    """Extracts skills and removes skill level (since it's not in the response)."""
    return [skill for skill in data.get("Skills", [])]


def extract_work_experience(data):
    """Extracts work experience details and formats responsibilities."""
    work_experience_list = []
    for job in data.get("Work Experience", []):
        work_experience_list.append({
            "company": job.get("Company", "Unknown"),
            "job_title": job.get("Job Title", "Unknown"),
            "start_date": job.get("Start Date", "Unknown"),
            "end_date": job.get("End Date", "Unknown"),
            "responsibilities": "\n".join(job.get("Responsibilities", []))  # Convert list to a string
        })
    return work_experience_list

def extract_certifications(data):
    """Extracts certifications according to the model, handling both dictionaries and plain text."""
    certification_list = []

    for cert in data.get("Certifications", []):
        if isinstance(cert, dict):  # Case: Certification is a dictionary
            certification_list.append({
                "name": cert.get("Name", "Unknown"),
                "issued_by": cert.get("Issued By", ""),
                "year": str(cert.get("Year", "")) if cert.get("Year") else None
            })
        elif isinstance(cert, str):  # Case: Certification is a string
            certification_list.append({
                "name": cert,
                "issued_by": "",
                "year": None
            })

    return certification_list

def parse_resume(text):
    """Main function that calls other functions to extract structured data."""
    groq_response = query_groq_resume_extraction(text)
    print("The groq response is ", groq_response)
    cleaned_response = clean_groq_response(groq_response)

    return {
        "personal_info": extract_personal_info(cleaned_response),
        "education": extract_education(cleaned_response),
        "skills": extract_skills(cleaned_response),
        "work_experience": extract_work_experience(cleaned_response),
        "certifications": extract_certifications(cleaned_response)
    }



def extract_text_from_pdf(pdf_path):
    images = convert_from_path(pdf_path)
    text = ""
    for image in images:
        text += pytesseract.image_to_string(image) + "\n"
    return text

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    return ""
